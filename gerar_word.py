from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Dados
data = {
    "funcao": "Pintor + ruído",
    "exames": {
        "admissional": "Exame Clínico | Audiometria | Dermatológico | TGO / TGP/ GGT | Ureia | Creatinina | Espirometria | RX Tórax PA OIT | Acuidade visual",
        "periodico": "Exame Clínico | Audiometria | Dermatológico | TGO / TGP/ GGT | Ureia | Creatinina | Espirometria | RX Tórax PA OIT | Acuidade visual",
        "retorno": "Exame Clínico",
        "mudanca_risco": "Realizar os exames complementares da periodicidade 'Demissional' e acrescentar os exames correspondentes ao 'admissional' do novo GHE.",
        "demissional": "Exame Clínico | Audiometria | Hemograma completo | Espirometria | Raio X Tórax PA OIT"
    },
    "riscos": {
        "Acidentes": "Incêndios ou explosões, respingo de produtos químicos nos olhos, choques ou curto circuito, quedas e escorregões",
        "Ergonômicos": "Postura inadequada e lesões por esforço repetitivo",
        "Físicos": "Ruído, radiação não ionizante (sol)",
        "Biológicos": "N/A",
        "Químicos": "Solventes, pigmentos, resina, cargas e aditivos"
    }
}

# Criar documento
doc = Document()

# Adicionar título principal
titulo = doc.add_heading('Consulta de Exames por Função', 0)
titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Adicionar nome da função
paragrafo_funcao = doc.add_paragraph()
run_funcao = paragrafo_funcao.add_run(f"Função: {data['funcao']}")
run_funcao.bold = True
run_funcao.font.size = Pt(14)
paragrafo_funcao.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Adicionar um espaço antes da tabela
doc.add_paragraph("\n")

# Criar tabela de exames
doc.add_heading('Tabela de Exames', level=1)
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

# Estilizar cabeçalho da tabela
header_cells = table.rows[0].cells
headers = ["ADMISSIONAL", "PERIÓDICO", "RETORNO AO TRABALHO", "MUDANÇA DE RISCO", "DEMISSIONAL"]

for i, header in enumerate(headers):
    p = header_cells[i].paragraphs[0]
    p.add_run(header).bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Adicionar dados na tabela
row_cells = table.add_row().cells
exames = data["exames"]

for i, exame in enumerate(exames.values()):
    p = row_cells[i].paragraphs[0]
    p.add_run(exame).font.size = Pt(10)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Adicionar espaçamento
doc.add_paragraph("\n")

# Criar tabela de Perigos e Riscos (sem exames relacionados)
doc.add_heading('Tabela de Perigos e Riscos', level=1)
table_riscos = doc.add_table(rows=1, cols=1)  # Apenas uma coluna para os riscos
table_riscos.style = 'Table Grid'

# Cabeçalho
header_cells = table_riscos.rows[0].cells
header_cells[0].text = "Perigos e Riscos"

# Deixar o cabeçalho em negrito e centralizado
for cell in header_cells:
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].bold = True

# Adicionar riscos na tabela
riscos = data["riscos"]

for tipo_risco, descricao_risco in riscos.items():
    row_cells = table_riscos.add_row().cells
    row_cells[0].text = f"{tipo_risco}: {descricao_risco}"

# Ajustar formatação da tabela de Perigos e Riscos
for row in table_riscos.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.runs[0].font.size = Pt(10)

# Adicionar espaçamento
doc.add_paragraph("\n")

# Salvar documento
doc.save('Consulta_de_Exames_e_Riscos.docx')